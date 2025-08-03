# Enhanced File Processing Pipeline
# Batch operations, async processing, and advanced document handling

import asyncio
import json
import time
from typing import Dict, List, Any, Optional, Tuple, Callable
from datetime import datetime, timedelta
from dataclasses import dataclass, asdict
from concurrent.futures import ThreadPoolExecutor, as_completed
import hashlib
import mimetypes
from pathlib import Path
import io
import logging

# Document processing libraries
import pdfplumber
from PIL import Image
import docx
import openpyxl
import csv

from core import log_debug, track_function_entry, bucket, index_endpoint, global_state
from core.ai.intelligence.ai_service import split_text, embed_text
from shared.config.base_config import DEPLOYED_INDEX_ID
from core.data.storage.cache_service import cache_service

@dataclass
class ProcessingResult:
    """Result of document processing"""
    file_path: str
    status: str  # success, error, skipped
    chunks_created: int
    processing_time_ms: float
    file_size_bytes: int
    error_message: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None

@dataclass
class BatchProcessingStats:
    """Statistics for batch processing operation"""
    total_files: int
    successful: int
    failed: int
    skipped: int
    total_chunks: int
    total_processing_time_ms: float
    start_time: datetime
    end_time: Optional[datetime] = None

class EnhancedDocumentProcessor:
    """Advanced document processor with multiple format support"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.md': self._process_text,
            '.docx': self._process_docx,
            '.xlsx': self._process_excel,
            '.csv': self._process_csv,
            '.json': self._process_json
        }
        
        # Processing configuration
        self.max_file_size_mb = 50
        self.chunk_size = 1000
        self.chunk_overlap = 200
        
        self.executor = ThreadPoolExecutor(max_workers=4)
        
        log_debug("Enhanced document processor initialized", {
            "supported_formats": list(self.supported_formats.keys()),
            "max_workers": 4
        })
    
    async def process_single_file(self, file_path: str, file_content: bytes, 
                                 file_metadata: Dict[str, Any] = None) -> ProcessingResult:
        """Process a single file with comprehensive error handling"""
        track_function_entry("process_single_file")
        start_time = time.time()
        
        try:
            # Validate file size
            file_size = len(file_content)
            if file_size > self.max_file_size_mb * 1024 * 1024:
                return ProcessingResult(
                    file_path=file_path,
                    status="error",
                    chunks_created=0,
                    processing_time_ms=0,
                    file_size_bytes=file_size,
                    error_message=f"File too large: {file_size / (1024*1024):.1f}MB > {self.max_file_size_mb}MB"
                )
            
            # Detect file format
            file_extension = Path(file_path).suffix.lower()
            if file_extension not in self.supported_formats:
                return ProcessingResult(
                    file_path=file_path,
                    status="error",
                    chunks_created=0,
                    processing_time_ms=0,
                    file_size_bytes=file_size,
                    error_message=f"Unsupported file format: {file_extension}"
                )
            
            # Check cache for existing processing result
            file_hash = hashlib.md5(file_content).hexdigest()
            cached_result = cache_service.get_document_metadata(f"{file_path}:{file_hash}")
            if cached_result:
                log_debug("Using cached processing result", {"file": file_path})
                return ProcessingResult(**cached_result)
            
            # Process file content
            processor_func = self.supported_formats[file_extension]
            text_content = await self._run_in_executor(processor_func, file_content, file_path)
            
            if not text_content or not text_content.strip():
                return ProcessingResult(
                    file_path=file_path,
                    status="error",
                    chunks_created=0,
                    processing_time_ms=(time.time() - start_time) * 1000,
                    file_size_bytes=file_size,
                    error_message="No text content extracted"
                )
            
            # Create chunks
            chunks = split_text(text_content)
            if not chunks:
                return ProcessingResult(
                    file_path=file_path,
                    status="error",
                    chunks_created=0,
                    processing_time_ms=(time.time() - start_time) * 1000,
                    file_size_bytes=file_size,
                    error_message="Could not create text chunks"
                )
            
            # Process chunks and create embeddings
            chunks_created = await self._process_chunks(file_path, chunks)
            
            processing_time_ms = (time.time() - start_time) * 1000
            
            # Create result
            result = ProcessingResult(
                file_path=file_path,
                status="success",
                chunks_created=chunks_created,
                processing_time_ms=processing_time_ms,
                file_size_bytes=file_size,
                metadata={
                    "file_hash": file_hash,
                    "text_length": len(text_content),
                    "chunk_count": len(chunks),
                    "file_extension": file_extension,
                    "processing_timestamp": datetime.utcnow().isoformat()
                }
            )
            
            # Cache result
            cache_service.cache_document_metadata(f"{file_path}:{file_hash}", asdict(result))
            
            log_debug("File processed successfully", {
                "file": file_path,
                "chunks": chunks_created,
                "time_ms": processing_time_ms
            })
            
            return result
            
        except Exception as e:
            processing_time_ms = (time.time() - start_time) * 1000
            log_debug("File processing failed", {
                "file": file_path,
                "error": str(e),
                "time_ms": processing_time_ms
            })
            
            return ProcessingResult(
                file_path=file_path,
                status="error",
                chunks_created=0,
                processing_time_ms=processing_time_ms,
                file_size_bytes=len(file_content),
                error_message=str(e)
            )
    
    async def process_batch(self, files: List[Tuple[str, bytes, Dict[str, Any]]], 
                           progress_callback: Optional[Callable] = None) -> BatchProcessingStats:
        """Process multiple files in batch with progress tracking"""
        track_function_entry("process_batch")
        
        stats = BatchProcessingStats(
            total_files=len(files),
            successful=0,
            failed=0,
            skipped=0,
            total_chunks=0,
            total_processing_time_ms=0,
            start_time=datetime.utcnow()
        )
        
        log_debug("Starting batch processing", {"file_count": len(files)})
        
        # Process files concurrently
        tasks = []
        for file_path, file_content, file_metadata in files:
            task = asyncio.create_task(
                self.process_single_file(file_path, file_content, file_metadata)
            )
            tasks.append((task, file_path))
        
        # Wait for completion with progress tracking
        for i, (task, file_path) in enumerate(tasks):
            try:
                result = await task
                
                # Update statistics
                if result.status == "success":
                    stats.successful += 1
                    stats.total_chunks += result.chunks_created
                elif result.status == "error":
                    stats.failed += 1
                else:
                    stats.skipped += 1
                
                stats.total_processing_time_ms += result.processing_time_ms
                
                # Call progress callback
                if progress_callback:
                    progress_callback({
                        "completed": i + 1,
                        "total": len(files),
                        "current_file": file_path,
                        "result": result,
                        "stats": stats
                    })
                
            except Exception as e:
                stats.failed += 1
                log_debug("Batch processing task failed", {
                    "file": file_path,
                    "error": str(e)
                })
        
        stats.end_time = datetime.utcnow()
        
        log_debug("Batch processing completed", {
            "total_files": stats.total_files,
            "successful": stats.successful,
            "failed": stats.failed,
            "total_chunks": stats.total_chunks,
            "total_time_ms": stats.total_processing_time_ms
        })
        
        return stats
    
    async def _run_in_executor(self, func: Callable, *args) -> Any:
        """Run CPU-intensive function in executor"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, func, *args)
    
    async def _process_chunks(self, file_path: str, chunks: List[str]) -> int:
        """Process chunks and create embeddings"""
        embeddings_to_upsert = []
        uploaded_chunks = []
        
        for i, chunk in enumerate(chunks):
            chunk_path = f"chunks/{file_path}/{i}.txt"
            
            # Store chunk in GCS
            if self._upload_chunk_to_gcs(chunk_path, chunk):
                uploaded_chunks.append(chunk_path)
                
                # Create embedding (with caching)
                vector = cache_service.get_embedding(chunk)
                if not vector:
                    vector = embed_text(chunk)
                    cache_service.cache_embedding(chunk, vector)
                
                # Prepare datapoint for vector index
                datapoint = {
                    "datapoint_id": chunk_path,
                    "feature_vector": vector,
                    "restricts": [{"namespace": "filepath", "allow_list": [file_path]}]
                }
                embeddings_to_upsert.append(datapoint)
        
        # Upsert to vector index
        if index_endpoint and embeddings_to_upsert:
            try:
                index_endpoint.upsert_datapoints(
                    deployed_index_id=DEPLOYED_INDEX_ID,
                    datapoints=embeddings_to_upsert
                )
                log_debug(f"Upserted {len(embeddings_to_upsert)} datapoints for {file_path}")
            except Exception as e:
                log_debug("Error upserting to index", {
                    "file": file_path,
                    "error": str(e)
                })
                return 0
        
        return len(uploaded_chunks)
    
    def _upload_chunk_to_gcs(self, chunk_path: str, chunk_content: str) -> bool:
        """Upload chunk to Google Cloud Storage"""
        try:
            if bucket:
                blob = bucket.blob(chunk_path)
                blob.upload_from_string(chunk_content, content_type='text/plain')
                return True
        except Exception as e:
            log_debug("Error uploading chunk", {
                "chunk_path": chunk_path,
                "error": str(e)
            })
        return False
    
    # File format processors
    def _process_pdf(self, file_content: bytes, filename: str) -> str:
        """Process PDF files"""
        try:
            texts = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        # Extract text
                        page_text = page.extract_text()
                        if page_text:
                            texts.append(page_text)
                        
                        # Extract tables
                        tables = page.extract_tables()
                        for table in tables:
                            if table:
                                try:
                                    table_md = "\n".join([
                                        " | ".join([str(cell) if cell else "" for cell in row]) 
                                        for row in table
                                    ])
                                    texts.append(f"\n[Table on page {page_num + 1}]\n{table_md}\n")
                                except Exception:
                                    continue
                    except Exception as e:
                        log_debug(f"Error processing PDF page {page_num}", {"error": str(e)})
                        continue
            
            return "\n\n".join(texts)
            
        except Exception as e:
            raise Exception(f"PDF processing failed: {str(e)}")
    
    def _process_text(self, file_content: bytes, filename: str) -> str:
        """Process text files (txt, md)"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all fail, use utf-8 with error handling
            return file_content.decode('utf-8', errors='ignore')
            
        except Exception as e:
            raise Exception(f"Text processing failed: {str(e)}")
    
    def _process_docx(self, file_content: bytes, filename: str) -> str:
        """Process Word documents"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            texts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    texts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
                
                if table_text:
                    texts.append("\n[Table]\n" + "\n".join(table_text) + "\n")
            
            return "\n\n".join(texts)
            
        except Exception as e:
            raise Exception(f"DOCX processing failed: {str(e)}")
    
    def _process_excel(self, file_content: bytes, filename: str) -> str:
        """Process Excel files"""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True)
            texts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = [f"\n[Sheet: {sheet_name}]\n"]
                
                # Process rows
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        sheet_text.append(row_text)
                
                if len(sheet_text) > 1:  # More than just the header
                    texts.extend(sheet_text)
            
            return "\n".join(texts)
            
        except Exception as e:
            raise Exception(f"Excel processing failed: {str(e)}")
    
    def _process_csv(self, file_content: bytes, filename: str) -> str:
        """Process CSV files"""
        try:
            # Detect encoding
            text_content = self._process_text(file_content, filename)
            
            # Parse CSV
            csv_reader = csv.reader(io.StringIO(text_content))
            rows = []
            
            for row in csv_reader:
                if row:  # Skip empty rows
                    rows.append(" | ".join(row))
            
            return "\n".join(rows)
            
        except Exception as e:
            raise Exception(f"CSV processing failed: {str(e)}")
    
    def _process_json(self, file_content: bytes, filename: str) -> str:
        """Process JSON files"""
        try:
            text_content = self._process_text(file_content, filename)
            json_data = json.loads(text_content)
            
            # Convert JSON to readable text
            def json_to_text(obj, prefix=""):
                if isinstance(obj, dict):
                    texts = []
                    for key, value in obj.items():
                        texts.append(f"{prefix}{key}: {json_to_text(value, prefix + '  ')}")
                    return "\n".join(texts)
                elif isinstance(obj, list):
                    texts = []
                    for i, item in enumerate(obj):
                        texts.append(f"{prefix}[{i}]: {json_to_text(item, prefix + '  ')}")
                    return "\n".join(texts)
                else:
                    return str(obj)
            
            return json_to_text(json_data)
            
        except Exception as e:
            raise Exception(f"JSON processing failed: {str(e)}")

class BatchProcessingManager:
    """Manages batch processing operations with queue and scheduling"""
    
    def __init__(self):
        self.processor = EnhancedDocumentProcessor()
        self.processing_queue = asyncio.Queue()
        self.active_batches: Dict[str, BatchProcessingStats] = {}
        self.batch_history: List[BatchProcessingStats] = []
        self.max_history = 50
        
        log_debug("Batch processing manager initialized")
    
    async def submit_batch(self, batch_id: str, files: List[Tuple[str, bytes, Dict[str, Any]]], 
                          progress_callback: Optional[Callable] = None) -> str:
        """Submit a batch processing job"""
        
        # Add to queue
        await self.processing_queue.put({
            "batch_id": batch_id,
            "files": files,
            "progress_callback": progress_callback,
            "submitted_at": datetime.utcnow()
        })
        
        log_debug("Batch submitted", {
            "batch_id": batch_id,
            "file_count": len(files)
        })
        
        return batch_id
    
    async def process_queue(self):
        """Process batches from queue"""
        while True:
            try:
                # Get next batch
                batch_job = await self.processing_queue.get()
                batch_id = batch_job["batch_id"]
                
                log_debug("Processing batch", {"batch_id": batch_id})
                
                # Process batch
                stats = await self.processor.process_batch(
                    batch_job["files"],
                    batch_job["progress_callback"]
                )
                
                # Store results
                self.active_batches[batch_id] = stats
                self.batch_history.append(stats)
                
                # Cleanup old history
                if len(self.batch_history) > self.max_history:
                    self.batch_history.pop(0)
                
                log_debug("Batch completed", {
                    "batch_id": batch_id,
                    "stats": asdict(stats)
                })
                
            except Exception as e:
                log_debug("Batch processing error", {"error": str(e)})
    
    def get_batch_status(self, batch_id: str) -> Optional[Dict[str, Any]]:
        """Get status of a batch"""
        if batch_id in self.active_batches:
            return asdict(self.active_batches[batch_id])
        return None
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get overall processing statistics"""
        if not self.batch_history:
            return {"status": "no_batches_processed"}
        
        total_files = sum(batch.total_files for batch in self.batch_history)
        total_successful = sum(batch.successful for batch in self.batch_history)
        total_failed = sum(batch.failed for batch in self.batch_history)
        total_chunks = sum(batch.total_chunks for batch in self.batch_history)
        
        return {
            "total_batches": len(self.batch_history),
            "total_files_processed": total_files,
            "success_rate": (total_successful / total_files * 100) if total_files > 0 else 0,
            "total_chunks_created": total_chunks,
            "total_failed": total_failed,
            "recent_batches": [asdict(batch) for batch in self.batch_history[-5:]]
        }

# Global instances
enhanced_processor = EnhancedDocumentProcessor()
batch_manager = BatchProcessingManager()